import os
import re
import time
import logging
import asyncio
from datetime import datetime, timezone
from typing import TypedDict, Any, List, Tuple, Dict, Optional

import chromadb
from openai import AsyncOpenAI, APITimeoutError, APIStatusError
from rank_bm25 import BM25Okapi
from tenacity import (
    retry,
    stop_after_attempt,
    wait_exponential,
    retry_if_exception_type,
    before_sleep_log,
)
from langchain_openai import ChatOpenAI
from langchain_core.messages import SystemMessage, HumanMessage
from langgraph.graph import StateGraph, END

logger = logging.getLogger(__name__)

NVIDIA_BASE_URL = "https://integrate.api.nvidia.com/v1"
EMBED_MODEL     = "nvidia/nv-embedqa-e5-v5"
LLM_MODEL       = "openai/gpt-oss-20b"

CHUNK_SIZE           = 600
CHUNK_OVERLAP        = 120
MIN_CHUNK_CHARS      = 100
MAX_CHUNKS_PER_DOC   = 1000
PENDING_TTL_SECONDS  = 3600
MIN_RELEVANCE_SCORE  = 0.10

# ── Text extraction ───────────────────────────────────────────────────────────

def extract_txt(content: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return content.decode(enc)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")

def extract_pdf(content: bytes) -> str:
    import io
    try:
        import pdfplumber
        parts =[]
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
        return "\n\n".join(parts)
    except ImportError:
        pass
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(content))
    return "\n\n".join([page.extract_text() for page in reader.pages if page.extract_text()])

def extract_docx(content: bytes) -> str:
    import io
    from docx import Document
    doc = Document(io.BytesIO(content))
    parts =[p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)

# ── Chunking ──────────────────────────────────────────────────────────────────

def chunk_text(text: str, chunk_tokens: int = CHUNK_SIZE, overlap_tokens: int = CHUNK_OVERLAP) -> list[str]:
    chunk_chars   = chunk_tokens * 4
    overlap_chars = overlap_tokens * 4

    text = re.sub(r"\n{3,}", "\n\n", text)
    paragraphs =[p.strip() for p in text.split("\n\n") if p.strip()]

    chunks: list[str] =[]
    current = ""

    for para in paragraphs:
        if len(para) > chunk_chars:
            sentences =[s.strip() for s in para.split("\n") if s.strip()]
            for sentence in sentences:
                if len(sentence) > chunk_chars:
                    for i in range(0, len(sentence), chunk_chars - overlap_chars):
                        part = sentence[i : i + chunk_chars]
                        if len(current) + len(part) > chunk_chars and current:
                            chunks.append(current)
                            current = part
                        else:
                            current = (current + " " + part).strip()
                    continue
                if len(current) + len(sentence) + 2 <= chunk_chars:
                    current = (current + "\n" + sentence).strip()
                else:
                    if current and len(current) >= MIN_CHUNK_CHARS:
                        chunks.append(current)
                    overlap_text = current[-overlap_chars:] if len(current) > overlap_chars else current
                    current = (overlap_text + "\n" + sentence).strip()
        else:
            if len(current) + len(para) + 2 <= chunk_chars:
                current = (current + "\n\n" + para).strip()
            else:
                if current and len(current) >= MIN_CHUNK_CHARS:
                    chunks.append(current)
                overlap_text = current[-overlap_chars:] if len(current) > overlap_chars else current
                current = (overlap_text + "\n\n" + para).strip()

    if current and len(current) >= MIN_CHUNK_CHARS:
        chunks.append(current)

    merged: list[str] =[]
    for c in chunks:
        if merged and len(c) < MIN_CHUNK_CHARS:
            merged[-1] += "\n\n" + c
        else:
            merged.append(c)

    return merged

# ── LangGraph State ───────────────────────────────────────────────────────────

class RAGState(TypedDict):
    question:       str
    top_k:          int
    chunks:         list[tuple[str, dict, float, str]]  # doc_text, metadata, distance, source ("vector"|"bm25")
    answer:         str
    confidence:     float
    sources:        list[dict]
    tokens_used:    int | None
    retrieval_type: str | None

# ── RAGService ────────────────────────────────────────────────────────────────

class RAGService:
    def __init__(self):
        api_key = os.environ.get("NVIDIA_API_KEY", "")
        if not api_key:
            logger.warning("NVIDIA_API_KEY is not set. Using placeholder.")

        self._openai = AsyncOpenAI(
            base_url=NVIDIA_BASE_URL,
            api_key=api_key or "placeholder",
        )

        self._llm = ChatOpenAI(
            model=LLM_MODEL,
            base_url=NVIDIA_BASE_URL,
            api_key=api_key or "placeholder",
            temperature=0.2,
            max_tokens=3000,
        )

        self.chroma = chromadb.PersistentClient(
            path="./chroma_data",
            settings=chromadb.Settings(anonymized_telemetry=False),
        )
        self.collection = self.chroma.get_or_create_collection(
            name="documents",
            metadata={"hnsw:space": "cosine"},
        )

        self._pending:      dict[str, dict]         = {}
        self._doc_registry: dict[str, dict]         = {}
        self._index_locks:  dict[str, asyncio.Lock] = {}

        self.bm25: Optional[BM25Okapi] = None
        self.bm25_ids: List[str] =[]
        
        # Глобальный лимит на одновременные запросы к API эмбеддингов
        self._embed_sem = asyncio.Semaphore(5)

        self._build_bm25_sync()
        self._restore_doc_registry()
        self._graph = self._build_graph()

    def _build_graph(self):
        g = StateGraph(RAGState)
        g.add_node("retrieve",  self._retrieve_node)
        g.add_node("grade",     self._grade_node)
        g.add_node("generate",  self._generate_node)

        g.set_entry_point("retrieve")
        g.add_edge("retrieve", "grade")
        g.add_edge("grade",    "generate")
        g.add_edge("generate", END)

        return g.compile()

    async def _retrieve_node(self, state: RAGState) -> dict:
        question = state["question"]
        top_k    = state["top_k"]
        k_pool   = max(top_k * 3, 20)

        vector_ranks: dict[str, int] = {}
        bm25_ranks:   dict[str, int] = {}
        chunk_map:    dict[str, tuple[str, dict, float, str]] = {}

        # 1. Vector search
        q_emb = await self._embed([question], input_type="query")
        vec = await asyncio.to_thread(
            self.collection.query,
            query_embeddings=q_emb,
            n_results=min(k_pool, max(1, self.collection.count())),
            include=["documents", "metadatas", "distances"],
        )
        
        if vec["ids"] and vec["ids"][0]:
            for rank, (cid, doc, meta, dist) in enumerate(zip(
                vec["ids"][0], vec["documents"][0],
                vec["metadatas"][0], vec["distances"][0],
            )):
                vector_ranks[cid] = rank + 1
                chunk_map[cid]    = (doc, meta, dist, "vector")

        # 2. BM25 search
        if self.bm25:
            tokenized_q = self._tokenize(question)
            scores = await asyncio.to_thread(self.bm25.get_scores, tokenized_q)
            
            top_indices = sorted(range(len(scores)), key=lambda i: scores[i], reverse=True)[:k_pool]
            bm25_missing_ids =[]
            
            for rank, idx in enumerate(top_indices):
                if scores[idx] > 0.0:
                    cid = self.bm25_ids[idx]
                    bm25_ranks[cid] = rank + 1
                    if cid not in chunk_map:
                        bm25_missing_ids.append(cid)

            # Запрашиваем тексты из ChromaDB только для тех BM25-чанков, которых не было в Vector
            if bm25_missing_ids:
                bm25_docs = await asyncio.to_thread(
                    self.collection.get,
                    ids=bm25_missing_ids,
                    include=["documents", "metadatas"]
                )
                for cid, doc, meta in zip(bm25_docs["ids"], bm25_docs["documents"], bm25_docs["metadatas"]):
                    chunk_map[cid] = (doc, meta, 0.0, "bm25") 

        # 3. RRF (Reciprocal Rank Fusion)
        all_ids = set(vector_ranks) | set(bm25_ranks)
        rrf = {
            cid: (1.0 / (60 + vector_ranks.get(cid, 1000))) + (1.0 / (60 + bm25_ranks.get(cid, 1000)))
            for cid in all_ids
        }
        sorted_ids = sorted(rrf, key=lambda c: rrf[c], reverse=True)[:top_k]
        chunks = [chunk_map[cid] for cid in sorted_ids]

        logger.info(f"[retrieve] question='{question}' → {len(chunks)} chunks")
        return {"chunks": chunks}

    async def _grade_node(self, state: RAGState) -> dict:
        filtered =[
            (doc, meta, dist, source)
            for doc, meta, dist, source in state["chunks"]
            if source == "bm25" or (1 - dist) >= MIN_RELEVANCE_SCORE
        ]
        dropped = len(state["chunks"]) - len(filtered)
        if dropped:
            logger.info(f"[grade] dropped {dropped} low-relevance chunks")
        return {"chunks": filtered}

    async def _generate_node(self, state: RAGState) -> dict:
        chunks   = state["chunks"]
        question = state["question"]

        if not chunks:
            return {
                "answer":         "Не удалось найти релевантную информацию в загруженных документах.",
                "confidence":     0.0,
                "sources":[],
                "tokens_used":    0,
                "retrieval_type": None,
            }

        context = "\n\n---\n\n".join(
            f"[Источник: {m['file_name']}, чанк {m['chunk_index']+1}/{m['total_chunks']}]\n{d}"
            for d, m, _, _ in chunks
        )

        messages =[
            SystemMessage(content=(
                "Ты — точный аналитический ассистент. "
                "Отвечай строго на основе предоставленного контекста. "
                "Не придумывай факты, которых нет в контексте. "
                "Если информации недостаточно — скажи об этом прямо. "
                "Отвечай на том же языке, на котором задан вопрос. "
                "Структурируй ответ: краткий вывод, затем детали при необходимости."
            )),
            HumanMessage(content=f"Контекст:\n\n{context}\n\n---\n\nВопрос: {question}"),
        ]

        response    = await self._llm.ainvoke(messages)
        answer      = response.content
        tokens_used = response.usage_metadata.get("total_tokens") if response.usage_metadata else None

        def calibrate(sim: float) -> float:
            return max(0.01, min((sim - 0.15) / 0.25, 1.0))

        # Честная оценка уверенности и типа поиска
        vector_dists =[d for _, _, d, src in chunks if src == "vector"]
        bm25_count   = sum(1 for _, _, _, src in chunks if src == "bm25")

        if not vector_dists:
            # Нашли только по тексту (BM25), вектор не дал релевантных совпадений
            confidence = 0.40  
            retrieval_type = "bm25_only"
        else:
            max_sim = max(1 - d for d in vector_dists)
            confidence = round(calibrate(max_sim), 2)
            retrieval_type = "hybrid" if bm25_count > 0 else "vector_only"

        sources =[
            {
                "chunk_id":    f"{m['file_id']}_chunk_{m['chunk_index']:04d}",
                "file_name":   m["file_name"],
                "chunk_index": m["chunk_index"],
                "score":       round(calibrate(1 - dist), 2) if src == "vector" else 0.50,
                "source_type": src,
                "preview":     doc[:200] + ("…" if len(doc) > 200 else ""),
            }
            for doc, m, dist, src in chunks[:5]
        ]

        logger.info(f"[generate] type={retrieval_type}, conf={confidence}, tokens={tokens_used}")
        return {
            "answer":         answer,
            "confidence":     confidence,
            "sources":        sources,
            "tokens_used":    tokens_used,
            "retrieval_type": retrieval_type,
        }

    # ── BM25 ──────────────────────────────────────────────────────────────────

    def _tokenize(self, text: str) -> list[str]:
        return[w for w in re.split(r"\W+", text.lower()) if w]

    def _build_bm25_sync(self):
        if self.collection.count() == 0:
            self.bm25 = None
            self.bm25_ids =[]
            return
        
        docs = self.collection.get(include=["documents"])
        self.bm25_ids = docs["ids"]
        tokenized = [self._tokenize(d) for d in docs["documents"]]
        
        if tokenized:
            self.bm25 = BM25Okapi(tokenized)
            logger.info(f"BM25 built: {len(self.bm25_ids)} chunks")

    async def rebuild_bm25(self):
        await asyncio.to_thread(self._build_bm25_sync)

    # ── Doc registry ──────────────────────────────────────────────────────────

    def _restore_doc_registry(self):
        if self.collection.count() == 0:
            return
        results = self.collection.get(include=["metadatas"])
        for meta in results["metadatas"]:
            fid = meta.get("file_id", "")
            if fid and fid not in self._doc_registry:
                self._doc_registry[fid] = {
                    "file_id":    fid,
                    "file_name":  meta.get("file_name", ""),
                    "file_type":  meta.get("file_type", ""),
                    "indexed_at": meta.get("indexed_at", ""),
                    "chunks":     0,
                }
            if fid:
                self._doc_registry[fid]["chunks"] += 1
        logger.info(f"Doc registry restored: {len(self._doc_registry)} docs")

    def _cleanup_expired_pending(self):
        now     = time.time()
        expired =[
            fid for fid, info in self._pending.items()
            if now - info["uploaded_ts"] > PENDING_TTL_SECONDS
        ]
        for fid in expired:
            logger.warning(f"Pending expired: {self._pending[fid]['file_name']} ({fid})")
            del self._pending[fid]

    # ── Public helpers ────────────────────────────────────────────────────────

    def extract_text(self, content: bytes, ext: str, filename: str) -> str:
        if ext in (".txt", ".md"):   return extract_txt(content)
        if ext == ".pdf":            return extract_pdf(content)
        if ext == ".docx":           return extract_docx(content)
        raise ValueError(f"Unsupported extension: {ext}")

    def store_pending(self, file_id: str, filename: str, ext: str, text: str):
        self._cleanup_expired_pending()
        self._pending[file_id] = {
            "file_id":     file_id,
            "file_name":   filename,
            "file_type":   ext,
            "text":        text,
            "uploaded_at": datetime.now(timezone.utc).isoformat(),
            "uploaded_ts": time.time(),
        }

    def total_chunks(self) -> int:
        return self.collection.count()

    def list_documents(self) -> list[dict]:
        return list(self._doc_registry.values())

    async def delete_document(self, file_id: str):
        results = await asyncio.to_thread(self.collection.get, where={"file_id": file_id})
        if results["ids"]:
            await asyncio.to_thread(self.collection.delete, ids=results["ids"])
            self._doc_registry.pop(file_id, None)
            await self.rebuild_bm25()

    # ── Indexing ──────────────────────────────────────────────────────────────

    def _get_index_lock(self, file_id: str) -> asyncio.Lock:
        if file_id not in self._index_locks:
            self._index_locks[file_id] = asyncio.Lock()
        return self._index_locks[file_id]

    async def index_file(self, file_id: str) -> dict:
        async with self._get_index_lock(file_id):
            if file_id not in self._pending:
                raise ValueError(f"File {file_id} not found in pending queue")

            info     = self._pending[file_id]
            text     = info["text"]
            filename = info["file_name"]

            old = await asyncio.to_thread(self.collection.get, where={"file_id": file_id})
            if old["ids"]:
                await asyncio.to_thread(self.collection.delete, ids=old["ids"])

            chunks = await asyncio.to_thread(chunk_text, text)
            if not chunks:
                raise ValueError("No chunks produced from document")
            
            if len(chunks) > MAX_CHUNKS_PER_DOC:
                pct = 100 - round(MAX_CHUNKS_PER_DOC / len(chunks) * 100)
                logger.warning(
                    f"'{filename}': {len(chunks)} chunks → truncated to {MAX_CHUNKS_PER_DOC} "
                    f"(~{pct}% of content ignored)"
                )
                chunks = chunks[:MAX_CHUNKS_PER_DOC]

            indexed_at = datetime.now(timezone.utc).isoformat()
            
            async def get_embedding_batch(batch_chunks: list[str]) -> list[list[float]]:
                # Используем глобальный семафор класса
                async with self._embed_sem:
                    return await self._embed(batch_chunks, input_type="passage")

            tasks =[
                get_embedding_batch(chunks[i : i + 32])
                for i in range(0, len(chunks), 32)
            ]
            
            batch_results = await asyncio.gather(*tasks)
            all_embeddings = [emb for batch in batch_results for emb in batch]

            ids       =[f"{file_id}_chunk_{i:04d}" for i in range(len(chunks))]
            metadatas =[
                {
                    "file_id":      file_id,
                    "file_name":    filename,
                    "file_type":    info["file_type"],
                    "chunk_index":  i,
                    "total_chunks": len(chunks),
                    "indexed_at":   indexed_at,
                }
                for i in range(len(chunks))
            ]

            await asyncio.to_thread(
                self.collection.add,
                ids=ids, embeddings=all_embeddings,
                documents=chunks, metadatas=metadatas,
            )

            self._doc_registry[file_id] = {
                "file_id":    file_id,
                "file_name":  filename,
                "file_type":  info["file_type"],
                "indexed_at": indexed_at,
                "chunks":     len(chunks),
            }

            del self._pending[file_id]
            self._index_locks.pop(file_id, None)

            # Вызов rebuild_bm25 удалён, он выполняется в main.py после батча
            logger.info(f"Indexed '{filename}': {len(chunks)} chunks")

            return {
                "file_id":   file_id,
                "file_name": filename,
                "chunks":    len(chunks),
                "status":    "indexed",
            }

    async def query(self, question: str, top_k: int = 7) -> dict:
        initial_state: RAGState = {
            "question":       question,
            "top_k":          top_k,
            "chunks":[],
            "answer":         "",
            "confidence":     0.0,
            "sources":[],
            "tokens_used":    None,
            "retrieval_type": None,
        }
        result = await self._graph.ainvoke(initial_state)
        return {
            "answer":         result["answer"],
            "confidence":     result["confidence"],
            "sources":        result["sources"],
            "tokens_used":    result["tokens_used"],
            "retrieval_type": result.get("retrieval_type"),
        }

    @retry(
        retry=retry_if_exception_type((APITimeoutError, APIStatusError)),
        wait=wait_exponential(multiplier=1, min=2, max=30),
        stop=stop_after_attempt(4),
        before_sleep=before_sleep_log(logger, logging.WARNING),
        reraise=True,
    )
    async def _embed(self, texts: list[str], input_type: str = "passage") -> list[list[float]]:
        response = await self._openai.embeddings.create(
            model=EMBED_MODEL,
            input=texts,
            extra_body={"input_type": input_type, "truncate": "END"},
        )
        return[item.embedding for item in response.data]