import re
import io
from langchain_text_splitters import RecursiveCharacterTextSplitter
from .config import (
    CHARS_PER_TOKEN,
    CHUNK_SIZE,
    CHUNK_OVERLAP,
    MIN_CHUNK_CHARS,
)

def count_tokens(text: str) -> int:
    return int(len(text) / CHARS_PER_TOKEN)

# ── Text extraction ───────────────────────────────────────────────────────────

def extract_txt(content: bytes) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            return content.decode(enc)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="replace")

def extract_pdf(content: bytes) -> str:
    try:
        import pdfplumber
        parts =[]
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    parts.append(t)
        return "\n\n".join(parts)
    except ImportError:
        pass
    
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(content))
    parts =[]
    for page in reader.pages:
        text = page.extract_text()
        if text:
            parts.append(text)
    return "\n\n".join(parts)

def extract_docx(content: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(content))
    parts =[p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)

def extract_text(content: bytes, ext: str, filename: str) -> str:
    if ext in (".txt", ".md"):   return extract_txt(content)
    if ext == ".pdf":            return extract_pdf(content)
    if ext == ".docx":           return extract_docx(content)
    raise ValueError(f"Unsupported extension: {ext}")


# ── Chunking (Улучшенный семантический сплиттер) ─────────────────────────────

def chunk_text(text: str, chunk_tokens: int = CHUNK_SIZE, overlap_tokens: int = CHUNK_OVERLAP) -> list[str]:
    # Очищаем слишком длинные пустоты
    text = re.sub(r"\n{3,}", "\n\n", text)
    
    chunk_chars   = chunk_tokens * 4
    overlap_chars = overlap_tokens * 4

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_chars,
        chunk_overlap=overlap_chars,
        separators=["\n\n", "\n", ".", "!", "?", " ", ""]
    )
    
    chunks = splitter.split_text(text)
    
    # Сливаем мелкие обрезки
    merged: list[str] =[]
    for c in chunks:
        if merged and len(c) < MIN_CHUNK_CHARS:
            merged[-1] += " " + c
        else:
            merged.append(c)

    return merged